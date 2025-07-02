from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()
doc.add_heading("III. PROPOSED METHODOLOGY 提出的方法", level=1)

sections = {
    "": """我们提出的藏文自然场景文本识别框架的总体架构如图 5 所示。
该框架基于五个组件：（1）残差型 CNN 特征提取器：采用 ResNet 风格连接增强训练稳定性；（2）双向 LSTM 序列建模器：在视觉与语言建模之间引入上下文增强层；（3）特征金字塔网络（FPN）：提取不同尺度的语义信息；（4）多维注意力机制模块：集成 CBAM 风格通道与空间注意力；（5）Transformer 编码器 + CTC 解码模块：进行全局建模与对齐。
整体流程如下：输入灰度图像 → Res-CNN → BiLSTM → FPN 特征融合 → 多维注意力增强 → Transformer → 输出字符概率 → CTC 损失训练。"""
,
    "A. CNN FOR FEATURE EXTRACTION": """为增强深层特征提取能力，我们使用 ResNet 风格的残差连接：y=f(x,{wi})+x，其中 f 为卷积、BN 和 ReLU 构成的残路径，x 为残差输入。此外，我们引入了特征金字塔结构（FPN）以融合不同层的语义信息，该策略提升了模型对不同尺度藏文字形的适应性。"""
,
    "B. BiLSTM + Transformer 编码器": """为了增强模型的序列建模能力，在 CNN 后我们添加双向 LSTM：
h_t = [→h_t; ←h_t]，其中 →h_t 和 ←h_t 分别表示前向和反向的隐藏状态，h_t 为双向合并后的输出。
Transformer 编码器由 N 层堆叠组成，每层包含多头自注意力机制和前馈子层：
Attention(Q,K,V)=softmax(QK^T/√d_k)V。
最终输出序列 Z ∈ RT×d，用于预测字符分布。"""
,
    "C. 多维注意力机制（CBAM + 序列注意力）": """使用 CBAM 结构对 CNN 输出加入注意力机制：
- 通道注意力：基于 max-pooling 和 avg-pooling 计算重要通道；
- 空间注意力：基于 2D 卷积判断图像中关键区域；
- 序列注意力：Transformer 进行跨步长自注意。
组合注意力提高模型在复杂背景和文字遮挡场景下的有效信息聚焦能力。
空间注意力的计算公式为：M_s = σ(Conv7×7([AvgPool; MaxPool]))，其中 σ 为 Sigmoid 激活函数。"""
,
    "D. CTC 解码与优化目标": """我们采用 CTC（Connectionist Temporal Classification）作为解码器与损失函数，其目标为：
L_ctc = −log∑_{π∈B(y)}P(π|x)，其中 B(y) 表示所有可映射到标签序列 y 的对齐路径，π_t, y_t 为预测概率。
最终损失为：L_CTC = ∑_{(x,y)∈D}L_ctc(x,y)。"""
,
    "E. 位置编码策略": """Transformer 编码器不具备顺序建模能力，我们引入正弦位置编码：
P_{pos,2i} = sin(pos/10000^{2i/d})，P_{pos,2i+1} = cos(pos/10000^{2i/d})。
其中 pos 是位置索引，d 是嵌入维度，该编码提供位置感知能力，辅助 Transformer 进行结构建模。"""
}

# Add content to document
for heading, text in sections.items():
    if heading:
        doc.add_heading(heading, level=2)
    doc.add_paragraph(text).style.font.size = Pt(12)

# Save document
file_path = "/mnt/data/Tibetan_Text_Recognition_Methodology.docx"
doc.save(file_path)

file_path
